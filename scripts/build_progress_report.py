from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TODO_PATH = ROOT / "todo.md"
OUTPUT = ROOT / "client_brief" / f"SAES_Client_Progress_Report_{date(2026, 6, 26).isoformat()}.docx"

BRAND_RED = RGBColor(0xE1, 0x2D, 0x3C)
BRAND_RED_DARK = RGBColor(0xB5, 0x16, 0x23)
INK = RGBColor(0x00, 0x34, 0x50)
BODY = RGBColor(0x19, 0x19, 0x19)
MUTED = RGBColor(0x5F, 0x66, 0x73)
BORDER = "D9DEE8"
SURFACE = "F4F4F4"
SUCCESS_FILL = "E8F3EC"
WARNING_FILL = "FFF4E5"
INFO_FILL = "F4F7FB"

SCREENSHOTS = [
    {
        "title": "Dashboard",
        "caption": "Operational overview with headline counts, alert surfaces, and links into the main workflows.",
        "path": ROOT / "client_brief" / "screenshots" / "dashboard.png",
    },
    {
        "title": "Assets",
        "caption": "Asset register with filtering, status handling, detail views, and QR-ready records.",
        "path": ROOT / "client_brief" / "screenshots" / "assets.png",
    },
    {
        "title": "Consumables",
        "caption": "Consumables and batch tracking screen for stock on hand, issue flows, and thresholds.",
        "path": ROOT / "client_brief" / "screenshots" / "consumables.png",
    },
    {
        "title": "Maintenance",
        "caption": "Maintenance schedules, service records, approved vendors, and due/overdue visibility.",
        "path": ROOT / "client_brief" / "screenshots" / "maintenance-2.png",
    },
    {
        "title": "Deployments",
        "caption": "Deployment planning and tracking with operational assignment flows for assets and consumables.",
        "path": ROOT / "client_brief" / "screenshots" / "deployments.png",
    },
    {
        "title": "Login",
        "caption": "Supabase-backed authentication screen now wired for real QA access.",
        "path": ROOT / "client_brief" / "screenshots" / "full-login-2.png",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_font(run, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Roboto"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Roboto")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Roboto")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def style_normal(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Roboto"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Roboto")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Roboto")
    normal.font.size = Pt(11)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("SAES Asset Register")
    set_font(run, 24, INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Client Progress Report")
    set_font(subtitle_run, 16, BRAND_RED, bold=True)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    meta_run = meta.add_run(
        f"Prepared for client review | {date(2026, 6, 26).strftime('%d %B %Y')}"
    )
    set_font(meta_run, 10.5, MUTED)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, 16, INK, bold=True)
    else:
        set_font(run, 12.5, BRAND_RED_DARK, bold=True)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_font(run, 11, BODY)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_font(run, 11, BODY)


def parse_todo_sections() -> list[dict[str, object]]:
    section_pattern = re.compile(r"^##\s+(\d+)\.\s+(.+)")
    item_pattern = re.compile(r"^- \[( |x)\] ")
    sections: list[dict[str, object]] = []
    current: dict[str, object] | None = None

    for line in TODO_PATH.read_text(encoding="utf-8").splitlines():
        section_match = section_pattern.match(line)
        if section_match:
            current = {
                "num": int(section_match.group(1)),
                "name": section_match.group(2).strip(),
                "done": 0,
                "total": 0,
            }
            sections.append(current)
            continue

        item_match = item_pattern.match(line)
        if item_match and current is not None:
            current["total"] = int(current["total"]) + 1
            if item_match.group(1) == "x":
                current["done"] = int(current["done"]) + 1

    return sections


def percent(done: int, total: int) -> float:
    if total == 0:
        return 0.0
    return round((done / total) * 100, 1)


def build_metrics():
    sections = parse_todo_sections()
    full_total = sum(int(section["total"]) for section in sections)
    full_done = sum(int(section["done"]) for section in sections)

    build_sections = [section for section in sections if 3 <= int(section["num"]) <= 25]
    build_total = sum(int(section["total"]) for section in build_sections)
    build_done = sum(int(section["done"]) for section in build_sections)

    core_sections = [section for section in sections if 3 <= int(section["num"]) <= 21]
    core_total = sum(int(section["total"]) for section in core_sections)
    core_done = sum(int(section["done"]) for section in core_sections)

    key_modules = [
        "Technical Foundation",
        "Supabase and Database Foundation",
        "Authentication and Roles",
        "Locations",
        "Asset Management",
        "Consumable Batches",
        "Stock Movements",
        "Planned Maintenance",
        "Deployments",
        "QR Codes and Scanning",
        "Attachments, Photos, and Documents",
        "Offline-First and Sync",
        "Dashboard",
        "Testing",
    ]
    module_rows = []
    for name in key_modules:
        match = next(section for section in sections if section["name"] == name)
        module_rows.append(
            {
                "name": name,
                "done": int(match["done"]),
                "total": int(match["total"]),
                "percent": percent(int(match["done"]), int(match["total"])),
            }
        )

    return {
        "full_done": full_done,
        "full_total": full_total,
        "full_percent": percent(full_done, full_total),
        "build_done": build_done,
        "build_total": build_total,
        "build_percent": percent(build_done, build_total),
        "core_done": core_done,
        "core_total": core_total,
        "core_percent": percent(core_done, core_total),
        "module_rows": module_rows,
    }


def status_label(module_percent: float) -> str:
    if module_percent >= 90:
        return "Substantially complete"
    if module_percent >= 75:
        return "Advanced"
    if module_percent >= 50:
        return "In progress"
    return "Early"


def status_fill(module_percent: float) -> str:
    if module_percent >= 90:
        return SUCCESS_FILL
    if module_percent >= 75:
        return INFO_FILL
    return WARNING_FILL


def add_summary_table(document: Document, metrics: dict[str, object]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [1.9, 1.25, 1.2, 2.15]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    headers = ["Measure", "Progress", "Status", "Notes"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, SURFACE)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, 10.5, INK, bold=True)

    rows = [
        (
            "Core application build",
            f"{metrics['core_percent']}%",
            "Preview-ready",
            "Main operational modules are in place and usable for guided review.",
        ),
        (
            "Implementation build",
            f"{metrics['build_percent']}%",
            "In progress",
            "Build work is well advanced, but testing and operational polish are still active.",
        ),
        (
            "Full project checklist",
            f"{metrics['full_percent']}%",
            "Mid-project",
            "Includes deployment, handoff, documentation, training, and broader readiness work.",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            if idx == 2:
                fill = INFO_FILL if text == "Preview-ready" else WARNING_FILL
                set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, 10.2, BODY, bold=(idx in {0, 1, 2}))


def add_module_table(document: Document, module_rows: list[dict[str, object]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [2.45, 1.0, 1.45, 1.6]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    headers = ["Module", "Percent", "Status", "Checklist"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, SURFACE)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, 10.2, INK, bold=True)

    for row in module_rows:
        cells = table.add_row().cells
        values = [
            str(row["name"]),
            f"{row['percent']}%",
            status_label(float(row["percent"])),
            f"{row['done']}/{row['total']}",
        ]
        for idx, text in enumerate(values):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            if idx == 2:
                set_cell_shading(cell, status_fill(float(row["percent"])))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, 10.0, BODY, bold=(idx == 0))


def add_callout(document: Document, heading: str, body: str, fill: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    set_cell_border(cell, color="D7DDE6")
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(heading)
    set_font(title_run, 10.5, INK, bold=True)

    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_run = body_p.add_run(body)
    set_font(body_run, 10.5, BODY)


def add_screenshot_page(document: Document, title: str, caption: str, path: Path) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(document, title, level=1)
    add_body(document, caption)

    if path.exists():
        document.add_picture(str(path), width=Inches(6.1))
        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Captured from the working preview environment.")
    set_font(run, 9.5, MUTED)


def build_report() -> Path:
    metrics = build_metrics()
    document = Document()
    style_normal(document)

    add_title_block(document)
    add_body(
        document,
        "This report summarises build progress for the Salvation Emergency Services Asset Register and provides a client-facing view of where the product stands today.",
    )

    add_heading(document, "Executive Summary", level=1)
    add_body(
        document,
        f"The core application build is now {metrics['core_percent']}% complete, covering login, assets, consumables, maintenance, deployments, attachments, dashboard, and offline foundations.",
    )
    add_body(
        document,
        f"Across the broader implementation build, the project is {metrics['build_percent']}% complete. Across the full end-to-end delivery checklist, which also includes documentation, deployment, training, and client rollout readiness, overall completion is {metrics['full_percent']}%.",
    )
    add_callout(
        document,
        "Testing readiness",
        "The application is now strong enough for a guided preview and internal QA, but it is not yet at the point I would classify as formal client testing readiness. The remaining work is mainly around audit trail completion, reporting/export, deeper test coverage, and deployment/handoff polish.",
        WARNING_FILL,
    )

    add_heading(document, "Progress Snapshot", level=1)
    add_summary_table(document, metrics)

    add_heading(document, "What Is Already Working", level=1)
    add_bullet(document, "Supabase-backed login with a working QA account for internal review.")
    add_bullet(document, "Asset management, including fleet/plant records and parent-child asset relationships.")
    add_bullet(document, "Consumable batch tracking, stock movement logic, and threshold alert foundations.")
    add_bullet(document, "Maintenance schedules, maintenance records, and approved vendor management.")
    add_bullet(document, "Deployment workflows linking assets and consumables to operational events.")
    add_bullet(document, "Attachments and document handling across operational records.")
    add_bullet(document, "Offline mutation queue and sync engine foundations.")
    add_bullet(document, "Operational dashboard with alert-oriented summary views.")

    add_heading(document, "Module Status", level=1)
    add_module_table(document, metrics["module_rows"])

    add_heading(document, "Before Formal Client Testing", level=1)
    add_bullet(document, "Finish the audit trail module and its related reporting surfaces.")
    add_bullet(document, "Complete reporting/export screens and validate the expected output set.")
    add_bullet(document, "Increase end-to-end and real-device QA, especially for scan and field workflows.")
    add_bullet(document, "Complete deployment and handoff readiness, including client-safe environment setup.")

    add_heading(document, "Preview Position", level=1)
    add_body(
        document,
        "The product has moved well beyond specification and scaffold stage. It is now a functional working build with the major operational modules visible and connected. My recommendation is to use the current build for a guided preview, not yet as a formal client test release.",
    )

    for screenshot in SCREENSHOTS:
        add_screenshot_page(document, screenshot["title"], screenshot["caption"], screenshot["path"])

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_report()
    print(output)
